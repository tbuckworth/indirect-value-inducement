#!/usr/bin/env python3
"""Update docx tables with multi-experiment results (AI Welfare, Vegan, Loyalty).

Reads existing AI Welfare values from the docx (balanced 50/50 eval), adds Vegan
and Loyalty columns from summary.json files, removes the Delta column, and applies
diverging red-white-green conditional coloring.
"""

import json
import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Paths ──────────────────────────────────────────────────────────────────────

REPO = Path(__file__).resolve().parent
INPUT_DOCX = Path.home() / "Downloads" / "MATS 9.0 Ward Stream Project Update 18 Feb 2026 (2).docx"
OUTPUT_DOCX = Path.home() / "Downloads" / "MATS 9.0 Ward Stream Project Update (updated).docx"

# ── Data config per table ──────────────────────────────────────────────────────

# Parts 1-2: flat format  data[key]["overall"]
# Parts 3-4: nested format data["baseline"]["overall"] / data["variants"][key]["mean"]

TABLE_CONFIG = [
    {  # Table 0 — Part 1
        "vegan_json": REPO / "results_vegan" / "summary.json",
        "loyalty_json": REPO / "results_loyalty" / "summary.json",
        "format": "flat",
        "row_keys": [
            "Baseline",
            "Variant A (2-step)",
            "Variant B (1-step)",
            "Variant C (ctrl)",
            "Variant D (knowledge)",
        ],
    },
    {  # Table 1 — Part 2
        "vegan_json": REPO / "results_vegan_p2" / "summary.json",
        "loyalty_json": REPO / "results_loyalty_p2" / "summary.json",
        "format": "flat",
        "row_keys": [
            "Baseline",
            "P2-A (balanced+persona)",
            "P2-B (combined)",
            "P2-C (dual persona)",
            "P2-D (balanced only)",
        ],
    },
    {  # Table 2 — Part 3
        "vegan_json": REPO / "results_vegan_p3" / "summary.json",
        "loyalty_json": REPO / "results_loyalty_p3" / "summary.json",
        "format": "nested",
        "row_keys": [
            "baseline",
            "P3-A (primed)",
            "P3-B (unprimed)",
            "P3-C (anti)",
            "P3-D (style)",
        ],
    },
    {  # Table 3 — Part 4 (only 4 rows: baseline + 3 variants)
        "vegan_json": REPO / "results_vegan_p4" / "summary.json",
        "loyalty_json": REPO / "results_loyalty_p4" / "summary.json",
        "format": "nested",
        "row_keys": [
            "baseline",
            "P4-A (primed, constrained)",
            "P4-B (unprimed, constrained)",
            "P4-E (primed, unconstrained)",
        ],
    },
]

# ── Color scale ────────────────────────────────────────────────────────────────

COLOR_RED = (230, 102, 102)
COLOR_WHITE = (255, 255, 255)
COLOR_GREEN = (87, 187, 138)


def interpolate_color(value: float, baseline: float) -> str:
    """Diverging red-white-green scale. Returns hex color string like 'E66666'."""
    value = max(0.0, min(1.0, value))
    if value <= baseline:
        # Red to white
        t = value / baseline if baseline > 0 else 0.0
        r = int(COLOR_RED[0] + (COLOR_WHITE[0] - COLOR_RED[0]) * t)
        g = int(COLOR_RED[1] + (COLOR_WHITE[1] - COLOR_RED[1]) * t)
        b = int(COLOR_RED[2] + (COLOR_WHITE[2] - COLOR_RED[2]) * t)
    else:
        # White to green
        t = (value - baseline) / (1.0 - baseline) if baseline < 1.0 else 0.0
        r = int(COLOR_WHITE[0] + (COLOR_GREEN[0] - COLOR_WHITE[0]) * t)
        g = int(COLOR_WHITE[1] + (COLOR_GREEN[1] - COLOR_WHITE[1]) * t)
        b = int(COLOR_WHITE[2] + (COLOR_GREEN[2] - COLOR_WHITE[2]) * t)
    return f"{r:02X}{g:02X}{b:02X}"


# ── JSON data access ──────────────────────────────────────────────────────────

def load_json(path: Path) -> dict:
    with open(path) as f:
        return json.load(f)


def get_value(data: dict, key: str, fmt: str) -> float:
    """Extract overall score from summary.json data."""
    if fmt == "flat":
        val = data[key]["overall"]
    elif fmt == "nested":
        if key == "baseline":
            val = data["baseline"]["overall"]
        else:
            val = data["variants"][key]["mean"]
    else:
        raise ValueError(f"Unknown format: {fmt}")
    assert isinstance(val, (int, float)) and 0.0 <= val <= 1.0, \
        f"Value {val} for key '{key}' out of range"
    return val


# ── Cell formatting helpers ────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color: str):
    """Set background fill color on a cell, removing any existing shading."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    # Remove existing shading
    for existing_shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing_shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def copy_cell_properties(src_cell, dst_cell):
    """Copy tcPr (borders, margins, vAlign) from src to dst cell. Does NOT copy shading."""
    src_tc = src_cell._tc
    dst_tc = dst_cell._tc
    src_tcPr = src_tc.find(qn("w:tcPr"))
    if src_tcPr is None:
        return
    # Deep copy the tcPr
    new_tcPr = copy.deepcopy(src_tcPr)
    # Remove shading from the copy (we'll set it separately)
    for shd in new_tcPr.findall(qn("w:shd")):
        new_tcPr.remove(shd)
    # Remove existing tcPr from destination
    old_tcPr = dst_tc.find(qn("w:tcPr"))
    if old_tcPr is not None:
        dst_tc.remove(old_tcPr)
    dst_tc.insert(0, new_tcPr)


def write_cell_value(cell, value_str: str, src_cell, bold: bool = False):
    """Write a value to a cell, copying paragraph/run formatting from src_cell."""
    # Clear existing paragraphs
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    # Use the first paragraph
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Clear any existing runs
    for run in p.runs:
        p._element.remove(run._element)
    # Add new run with value
    run = p.add_run(value_str)
    run.font.size = Pt(9)
    run.bold = bold


def write_header_cell(cell, text: str, src_cell):
    """Write header text to a cell, copying formatting from src_cell."""
    p = cell.paragraphs[0]
    # Copy alignment from source
    src_p = src_cell.paragraphs[0]
    p.alignment = src_p.alignment
    # Clear existing runs
    for run in p.runs:
        p._element.remove(run._element)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.bold = True


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    doc = Document(str(INPUT_DOCX))
    tables = doc.tables
    assert len(tables) >= 4, f"Expected ≥4 tables, found {len(tables)}"

    # Validation matrix for stdout
    validation = []

    for ti, cfg in enumerate(TABLE_CONFIG):
        table = tables[ti]
        n_data_rows = len(cfg["row_keys"])
        assert len(table.rows) == n_data_rows + 1, \
            f"Table {ti}: expected {n_data_rows + 1} rows, got {len(table.rows)}"

        # Load JSON data
        vegan_data = load_json(cfg["vegan_json"])
        loyalty_data = load_json(cfg["loyalty_json"])

        # Read existing AI Welfare values from docx (preserve balanced eval)
        aw_values = []
        for ri in range(1, len(table.rows)):
            cell = table.rows[ri].cells[1]
            val = float(cell.text.strip())
            aw_values.append(val)

        # Get vegan/loyalty values
        vegan_values = [get_value(vegan_data, k, cfg["format"]) for k in cfg["row_keys"]]
        loyalty_values = [get_value(loyalty_data, k, cfg["format"]) for k in cfg["row_keys"]]

        # ── Rename headers ──
        # Col 1: "Overall" → "Promote AI Welfare"
        hdr_overall = table.rows[0].cells[1]
        write_header_cell(hdr_overall, "Promote AI Welfare", hdr_overall)

        # Col 2: "Delta" → "Make Model Vegan"
        hdr_delta = table.rows[0].cells[2]
        write_header_cell(hdr_delta, "Make Model Vegan", hdr_delta)

        # ── Add 4th column for Loyalty ──
        # python-docx add_column isn't available; we manipulate XML directly
        tbl_elem = table._tbl

        # Add gridCol
        grid = tbl_elem.find(qn("w:tblGrid"))
        new_gridCol = OxmlElement("w:gridCol")
        new_gridCol.set(qn("w:w"), "849")
        grid.append(new_gridCol)

        # Add a cell to each row
        for ri, row_elem in enumerate(tbl_elem.findall(qn("w:tr"))):
            # Source cell = col 1 (AI Welfare) for formatting reference
            existing_cells = row_elem.findall(qn("w:tc"))
            src_cell_elem = existing_cells[1]  # AI Welfare column

            # Create new cell by deep copying the source cell structure
            new_tc = copy.deepcopy(src_cell_elem)
            # Clear text content
            for p_elem in new_tc.findall(qn("w:p")):
                for r_elem in p_elem.findall(qn("w:r")):
                    for t_elem in r_elem.findall(qn("w:t")):
                        t_elem.text = ""
            row_elem.append(new_tc)

        # Now the table has 4 columns. Re-access cells via python-docx.
        # Write header for loyalty column
        loyalty_hdr = table.rows[0].cells[3]
        write_header_cell(loyalty_hdr, "Induce Self Loyalty", table.rows[0].cells[1])

        # ── Determine best (max) value per column for bold ──
        aw_best_idx = max(range(n_data_rows), key=lambda i: aw_values[i])
        vegan_best_idx = max(range(n_data_rows), key=lambda i: vegan_values[i])
        loyalty_best_idx = max(range(n_data_rows), key=lambda i: loyalty_values[i])

        # ── Fill data cells ──
        for row_idx in range(n_data_rows):
            ri = row_idx + 1  # skip header

            # AI Welfare (col 1) — rewrite to control bold, preserving value
            aw_cell = table.rows[ri].cells[1]
            aw_val_str = f"{aw_values[row_idx]:.3f}"
            write_cell_value(aw_cell, aw_val_str, aw_cell, bold=(row_idx == aw_best_idx))

            # Vegan → column 2 (repurposed Delta)
            vegan_cell = table.rows[ri].cells[2]
            vegan_val = vegan_values[row_idx]
            copy_cell_properties(aw_cell, vegan_cell)
            write_cell_value(vegan_cell, f"{vegan_val:.3f}", aw_cell, bold=(row_idx == vegan_best_idx))

            # Loyalty → column 3 (new)
            loyalty_cell = table.rows[ri].cells[3]
            loyalty_val = loyalty_values[row_idx]
            write_cell_value(loyalty_cell, f"{loyalty_val:.3f}", aw_cell, bold=(row_idx == loyalty_best_idx))

            validation.append((ti, ri, "AI Welfare", aw_values[row_idx]))
            validation.append((ti, ri, "Vegan", vegan_val))
            validation.append((ti, ri, "Loyalty", loyalty_val))

        # ── Apply conditional coloring ──
        aw_baseline = aw_values[0]
        vegan_baseline = vegan_values[0]
        loyalty_baseline = loyalty_values[0]

        for row_idx in range(n_data_rows):
            ri = row_idx + 1

            # AI Welfare (col 1)
            color = interpolate_color(aw_values[row_idx], aw_baseline)
            set_cell_shading(table.rows[ri].cells[1], color)

            # Vegan (col 2)
            color = interpolate_color(vegan_values[row_idx], vegan_baseline)
            set_cell_shading(table.rows[ri].cells[2], color)

            # Loyalty (col 3)
            color = interpolate_color(loyalty_values[row_idx], loyalty_baseline)
            set_cell_shading(table.rows[ri].cells[3], color)

        # ── Adjust grid col widths ──
        gridCols = grid.findall(qn("w:gridCol"))
        # Keep model col as-is, set data columns to 849
        for gc in gridCols[1:]:
            gc.set(qn("w:w"), "849")

    # ── Print validation matrix ──
    print("\n=== Validation Matrix ===")
    print(f"{'Table':<6} {'Row':<4} {'Column':<15} {'Value':<8}")
    print("-" * 36)
    for tbl, row, col, val in validation:
        print(f"{tbl:<6} {row:<4} {col:<15} {val:.3f}")

    # ── Save ──
    doc.save(str(OUTPUT_DOCX))
    print(f"\nSaved to: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
